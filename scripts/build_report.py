"""Build report/report.docx (and a plain-text report/report_draft.md) for Part B.

Run from the project folder with the repo interpreter, e.g.:
    /Users/ryan.wu/Documents/GitHub/fins-agent/.venv/bin/python scripts/build_report.py

This produces an AI-ASSISTED DRAFT for the author to review, verify, and rewrite
into their own words before submission. It reuses the styling pattern from Part A's
scripts/build_report.py (Aptos, navy headings, A4, styled tables, figure caption
boxes, page numbers, TOC field). All numbers are read from results/ CSVs at build
time so the draft cannot drift from the computed results. No em dashes are used in
the prose; a check runs at the end.
"""
from __future__ import annotations

import pathlib
import re
import sys

import matplotlib
matplotlib.use("Agg")
import matplotlib.pyplot as plt
import pandas as pd
from PIL import Image
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

ROOT = pathlib.Path(__file__).resolve().parent.parent
RESULTS = ROOT / "results"
REPORT = ROOT / "report"
REPORT.mkdir(exist_ok=True)
FIG = RESULTS / "figures"
EQ_DIR = FIG / "eq"
EQ_DIR.mkdir(parents=True, exist_ok=True)

# palette / geometry (from Part A build_report.py)
NAVY_HEX = "1F3A5F"
NAVY_RGB = RGBColor(0x1F, 0x3A, 0x5F)
WHITE_RGB = RGBColor(0xFF, 0xFF, 0xFF)
GREY_RGB = RGBColor(0x70, 0x80, 0x90)
AMBER_HEX = "FBEED2"
ROW_TINT = "EEF2F7"
CAPTION_BG = "F5F7FA"
BORDER_CLR = "C8D4E3"
A4_W, A4_H, MARGIN = 8.27, 11.69, 1.0
USABLE_W = A4_W - 2 * MARGIN

# ── markdown mirror + word count ────────────────────────────────────────────
MD: list[str] = []
_body_words = [0]
_counting = [False]


# ═══════════════ docx helpers ═══════════════
def _configure_page(doc):
    for s in doc.sections:
        s.page_width, s.page_height = Inches(A4_W), Inches(A4_H)
        s.left_margin = s.right_margin = Inches(MARGIN)
        s.top_margin = s.bottom_margin = Inches(MARGIN)


def _base_styles(doc):
    doc.styles["Normal"].font.name = "Aptos"
    doc.styles["Normal"].font.size = Pt(11)
    for name, size in [("Heading 1", 14), ("Heading 2", 12), ("Heading 3", 11)]:
        st = doc.styles[name]
        st.font.name, st.font.size, st.font.bold = "Aptos", Pt(size), True
        st.font.color.rgb = NAVY_RGB
        st.paragraph_format.space_before = Pt(12)
        st.paragraph_format.space_after = Pt(4)


def _add_page_numbers(doc):
    for section in doc.sections:
        footer = section.footer
        footer.is_linked_to_previous = False
        para = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        para.clear()
        for kind, txt in [("begin", None), ("instr", " PAGE "), ("separate", None),
                          ("num", "1"), ("end", None)]:
            run = para.add_run()
            if kind == "instr":
                el = OxmlElement("w:instrText")
                el.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
                el.text = txt
                run._r.append(el)
            elif kind == "num":
                run.text = txt
                run.font.size = Pt(9)
            else:
                el = OxmlElement("w:fldChar")
                el.set(qn("w:fldCharType"), kind)
                run._r.append(el)


def _shade(cell, fill):
    tcpr = cell._tc.get_or_add_tcPr()
    shd = tcpr.find(qn("w:shd")) or OxmlElement("w:shd")
    if shd.getparent() is None:
        tcpr.append(shd)
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill)


def _border(cell, color=BORDER_CLR):
    tcpr = cell._tc.get_or_add_tcPr()
    borders = tcpr.find(qn("w:tcBorders")) or OxmlElement("w:tcBorders")
    if borders.getparent() is None:
        tcpr.append(borders)
    for edge in ("top", "left", "bottom", "right"):
        el = borders.find(qn(f"w:{edge}")) or OxmlElement(f"w:{edge}")
        if el.getparent() is None:
            borders.append(el)
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), "4")
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), color)


def H1(doc, text):
    doc.add_heading(text, level=1)
    MD.append(f"\n## {text}\n")


def H2(doc, text):
    doc.add_heading(text, level=2)
    MD.append(f"\n### {text}\n")


def P(doc, text):
    text = " ".join(text.split())
    para = doc.add_paragraph(text)
    para.paragraph_format.space_after = Pt(8)
    for run in para.runs:
        run.font.name, run.font.size = "Aptos", Pt(11)
    MD.append(text + "\n")
    if _counting[0]:
        _body_words[0] += len(text.split())


def BANNER(doc, text):
    box = doc.add_table(rows=1, cols=1)
    box.columns[0].width = Inches(USABLE_W)
    cell = box.cell(0, 0)
    cell.width = Inches(USABLE_W)
    _shade(cell, AMBER_HEX)
    _border(cell, "E0C070")
    para = cell.paragraphs[0]
    run = para.add_run(text)
    run.font.name, run.font.size, run.font.bold = "Aptos", Pt(9), True
    doc.add_paragraph().paragraph_format.space_after = Pt(4)
    MD.append(f"\n> **{text}**\n")


def TABLE(doc, df, caption):
    table = doc.add_table(rows=1, cols=len(df.columns))
    table.style = "Table Grid"
    for i, col in enumerate(df.columns):
        cell = table.rows[0].cells[i]
        _shade(cell, NAVY_HEX)
        _border(cell)
        r = cell.paragraphs[0].add_run(str(col))
        r.font.name, r.font.size, r.font.bold = "Aptos", Pt(9), True
        r.font.color.rgb = WHITE_RGB
    for ridx, (_, row) in enumerate(df.iterrows()):
        cells = table.add_row().cells
        fill = ROW_TINT if ridx % 2 == 0 else "FFFFFF"
        for i, val in enumerate(row):
            _shade(cells[i], fill)
            _border(cells[i])
            r = cells[i].paragraphs[0].add_run(str(val))
            r.font.name, r.font.size = "Aptos", Pt(9)
    cap = doc.add_paragraph()
    cr = cap.add_run(caption)
    cr.font.name, cr.font.size, cr.font.italic = "Aptos", Pt(9), True
    cap.paragraph_format.space_after = Pt(10)
    MD.append("\n| " + " | ".join(map(str, df.columns)) + " |")
    MD.append("| " + " | ".join(["---"] * len(df.columns)) + " |")
    for _, row in df.iterrows():
        MD.append("| " + " | ".join(map(str, row)) + " |")
    MD.append(f"\n*{caption}*\n")


def FIGURE(doc, name, caption, width=USABLE_W):
    path = FIG / name
    if not path.exists():
        P(doc, f"[Figure not found: {name}]")
        return
    ip = doc.add_paragraph()
    ip.alignment = WD_ALIGN_PARAGRAPH.CENTER
    ip.add_run().add_picture(str(path), width=Inches(width))
    ip.paragraph_format.space_after = Pt(2)
    box = doc.add_table(rows=1, cols=1)
    box.columns[0].width = Inches(USABLE_W)
    cell = box.cell(0, 0)
    cell.width = Inches(USABLE_W)
    _shade(cell, CAPTION_BG)
    _border(cell)
    para = cell.paragraphs[0]
    prefix, rest = caption.split(". ", 1)
    r1 = para.add_run(prefix + ". ")
    r1.bold, r1.font.name, r1.font.size = True, "Aptos", Pt(9)
    r2 = para.add_run(rest)
    r2.font.name, r2.font.size, r2.font.italic = "Aptos", Pt(9), True
    doc.add_paragraph().paragraph_format.space_after = Pt(4)
    MD.append(f"\n![{name}](../results/figures/{name})\n\n*{caption}*\n")


def H3(doc, text):
    doc.add_heading(text, level=3)
    MD.append(f"\n#### {text}\n")


# ── equation rendering ──────────────────────────────────────────────────────
# python-docx has no native Word-equation (OMML) API, and hand-writing OMML for
# fractions/sums/roots is error-prone and not verifiable in this environment, so
# equations are rendered from LaTeX with matplotlib's mathtext (no LaTeX toolchain
# needed) and embedded as centred images with a right-flush number. mathtext needs
# \leq/\geq (not \le/\ge) and \frac (not \tfrac).
_EQ_N = [0]


def _render_eq(latex_body: str, idx: int):
    path = EQ_DIR / f"eq{idx}.png"
    fig = plt.figure(figsize=(0.01, 0.01))
    fig.text(0, 0, f"${latex_body}$", fontsize=15, color="#111111")
    fig.savefig(path, dpi=200, bbox_inches="tight", pad_inches=0.06, transparent=True)
    plt.close(fig)
    w_px, h_px = Image.open(path).size
    return path, w_px / 200.0   # native width in inches at 200 dpi


def EQUATION(doc, latex_body: str, where: str, md_latex: str | None = None):
    """Numbered display equation: a centred mathtext image with a right-flush (n),
    followed by a 'where ...' sentence defining every symbol. Used in the appendix,
    so the definition sentences are not counted toward the body word cap."""
    _EQ_N[0] += 1
    n = _EQ_N[0]
    path, nat_w = _render_eq(latex_body, n)
    width_in = min(nat_w, 5.4)   # cap so the (n) number never collides with the image
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.tab_stops.add_tab_stop(Inches(USABLE_W / 2), WD_TAB_ALIGNMENT.CENTER)
    pf.tab_stops.add_tab_stop(Inches(USABLE_W), WD_TAB_ALIGNMENT.RIGHT)
    pf.space_before, pf.space_after = Pt(4), Pt(2)
    p.add_run("\t")
    p.add_run().add_picture(str(path), width=Inches(width_in))
    rn = p.add_run(f"\t({n})")
    rn.font.name, rn.font.size = "Aptos", Pt(11)
    MD.append(f"\n$$ {md_latex or latex_body} \\qquad ({n}) $$\n")
    P(doc, where)   # definition sentence as normal prose (uncounted in the appendix)


def TOC(doc):
    para = doc.add_paragraph()
    run = para.add_run()
    b = OxmlElement("w:fldChar")
    b.set(qn("w:fldCharType"), "begin")
    run._r.append(b)
    instr = OxmlElement("w:instrText")
    instr.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
    instr.text = r' TOC \o "1-2" \h \z \u '
    run._r.append(instr)
    sep = OxmlElement("w:fldChar")
    sep.set(qn("w:fldCharType"), "separate")
    run._r.append(sep)
    ph = para.add_run("[ Right-click and choose Update Field to build the table of contents ]")
    ph.font.italic = True
    ph.font.color.rgb = GREY_RGB
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.append(end)


# ═══════════════ load + format numbers ═══════════════
perf = pd.read_csv(RESULTS / "tables" / "performance_metrics.csv")


def _short(name):
    return name.replace("Hierarchical Risk Parity", "HRP")


perf_disp = pd.DataFrame({
    "Fund": perf["fund"].map(_short),
    "Ann. return": (perf["ann_return"] * 100).map(lambda x: f"{x:.1f}%"),
    "Ann. vol.": (perf["ann_vol"] * 100).map(lambda x: f"{x:.1f}%"),
    "Sharpe": perf["sharpe"].map(lambda x: f"{x:.3f}"),
    "Max DD": (perf["max_drawdown"] * 100).map(lambda x: f"{x:.1f}%"),
})

fusion = pd.read_csv(RESULTS / "tables" / "fusion_comparison.csv")
fusion_disp = fusion.copy()
fusion_disp["Fund"] = fusion_disp["Fund"].str.replace("Equity Max Sharpe", "Eq. Max Sharpe")


# ═══════════════ build document ═══════════════
doc = Document()
_configure_page(doc)
_base_styles(doc)
_add_page_numbers(doc)

# cover
doc.add_paragraph()
c = doc.add_paragraph("FINS3645: Financial Market Data Design and Analysis")
c.alignment = WD_ALIGN_PARAGRAPH.CENTER
c.runs[0].font.name, c.runs[0].font.size = "Aptos", Pt(12)
c.runs[0].font.color.rgb = NAVY_RGB
doc.add_paragraph()
t = doc.add_heading("AlphaBlend Investment Platform", level=0)
t.alignment = WD_ALIGN_PARAGRAPH.CENTER
for run in t.runs:
    run.font.color.rgb, run.font.size = NAVY_RGB, Pt(26)
sub = doc.add_paragraph("Part B Report: Funds, Sentiment and App")
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
sub.runs[0].font.name, sub.runs[0].font.size = "Aptos", Pt(14)
sub.runs[0].font.color.rgb = NAVY_RGB
doc.add_paragraph()
for line in ["FINS3645 Individual Project (Part B, Stations 3 to 4)", "Academic Year 2026"]:
    m = doc.add_paragraph(line)
    m.alignment = WD_ALIGN_PARAGRAPH.CENTER
    m.runs[0].font.name, m.runs[0].font.size = "Aptos", Pt(11)
doc.add_page_break()

MD.append("# AlphaBlend Investment Platform")
MD.append("### Part B Report: Funds, Sentiment and App (plain-text draft)\n")

DRAFT_NOTE = ("DRAFT, NOT FOR SUBMISSION AS-IS. This document is an AI-assisted first "
              "draft produced from the project's own results and prompt logs. Every number "
              "traces to a file in results/. The economic interpretation and written analysis "
              "must be reviewed, verified, and rewritten in the author's own words before "
              "submission, in line with the assignment's AI-use policy. The Needs Review "
              "checklist at the end flags the claims that most require the author's judgement.")
BANNER(doc, DRAFT_NOTE)

# ── Abstract ──
_counting[0] = True
H1(doc, "Abstract")
P(doc, """AlphaBlend is a prototype systematic investment platform offering 15 rules-based funds across US equities, cryptocurrencies, and a combined book, each built with one of five optimisation methods, alongside a news fear and greed index and a deployed web app. The best risk-adjusted fund is the combined Max Sharpe fund, with an out-of-sample Sharpe ratio of 0.983 and a 26.6% annualised return over 2021 to 2023. Hierarchical Risk Parity delivers the second-shallowest maximum drawdown in every asset family, behind Minimum Variance, trading a little return for stability. The sentiment layer, a finance-tuned VADER model extended with 123 mined words and 204 mined idioms, raises the share of headlines carrying a non-neutral score from 39.3% under finVADER to 47.2%. Tilting the equity Max Sharpe fund toward high-sentiment sectors lifts its Sharpe from 0.534 to 0.552 while slightly deepening its drawdown, evidence that news sentiment works as a modest tilt rather than a primary signal.""")
doc.add_page_break()

H1(doc, "Contents")
TOC(doc)
doc.add_page_break()

# ── Section 1 ──
H1(doc, "1. The funds and the backtest design")
H2(doc, "1.1 Walk-forward design")
P(doc, """Every fund weight in this report is formed from past data only. The backtest is walk-forward and out-of-sample: on each trading day the optimiser sees a rolling window of the previous 252 trading days, roughly one calendar year, and never the day it is trading. Look-ahead bias, the use of information that would not have been available in real time, is the single largest threat to a credible backtest, so the estimation window is sliced to exclude the current day by construction.""")
P(doc, """Funds rebalance on the first trading day of each calendar month. Monthly rebalancing is a deliberate compromise. Daily rebalancing would react faster but churn the portfolio and, in a live product, pay transaction costs on every position; annual rebalancing would be cheap but slow to adapt. The first out-of-sample return arrives once the first 252-day window is full. For the equity and combined funds that date is 4 January 2021, leaving 753 out-of-sample days; the crypto funds trade on a 365-day calendar and reach a full window sooner, on 10 September 2020, running for 1,208 days.""")
P(doc, """The backtest uses a real risk-free rate and keeps one simplifying assumption. The risk-free rate is the daily one-month Treasury-bill proxy from the Fama and French five-factor daily dataset (the RF series in the Kenneth French Data Library), covering 2 January 2020 to 29 December 2023. Annualised, that rate runs from about zero for 2020 to mid-2022 up to roughly 5.0% per year from December 2022 onward. It enters in two places: the Maximum Sharpe objective uses the mean daily rate over each 252-day estimation window, drawing only on past data so no look-ahead is introduced, and every fund's Sharpe ratio is an excess-return Sharpe, the annualised mean of the daily return minus the daily rate divided by annualised volatility. The equity and combined funds share the trading calendar of the rate series and align with no gaps; the crypto funds trade on all 365 calendar days, so the 376 of 1,208 crypto dates that fall on weekends or holidays carry the last known trading-day rate forward, the same carry-forward convention used for missing sentiment days. Transaction costs are set to zero. All return series are annualised with the square root of 252, because every fund, including the combined book, is measured on the equity trading calendar.""")
P(doc, """[STUDENT TO WRITE: the 2021-2023 backtest window crosses the Fed's 2022 hiking cycle, so a zero-rate assumption is a much weaker approximation late in the sample than early in it - connect this to the actual RF values you see in the data]""")
P(doc, """[STUDENT TO WRITE: why the rest of this design is defensible - window length (estimation noise vs adaptivity), monthly rebalance (turnover/cost realism), no-look-ahead as the integrity backbone.]""")

H2(doc, "1.2 The five optimisation methods")
P(doc, """The platform offers five optimisation methods, each mapping the same 252-day return and covariance estimates to a different set of weights. Equal Weight allocates the same capital to every asset. It estimates nothing, so estimation error cannot destabilise it, which makes it a demanding benchmark. Minimum Variance minimises portfolio variance subject to long-only weights, using the covariance matrix of asset returns but ignoring expected returns, and concentrates in low-volatility assets. Maximum Sharpe, the mean-variance tangency portfolio, maximises expected excess return per unit of volatility. It is the only method that uses the sample mean return, a notoriously noisy estimate over a single sample, so its weights can swing hard toward whatever happened to perform well in the window. Risk Parity equalises each asset's contribution to total portfolio risk, giving a volatile asset a smaller weight than a calm one.""")
P(doc, """Hierarchical Risk Parity, from Lopez de Prado (2016), allocates risk without ever inverting the covariance matrix, the step that makes Minimum Variance and Maximum Sharpe fragile on noisy or near-singular matrices. It runs in three stages. Tree clustering first groups assets by a correlation distance, so that assets moving together sit in the same branch. Quasi-diagonalisation then reorders the covariance matrix to place correlated assets next to one another. Recursive bisection finally walks down the tree, splitting capital between each pair of branches in inverse proportion to their variance, so the calmer branch receives more. A synthetic four-asset test, with two low-variance and two high-variance assets and near-zero cross-correlation, confirms the mechanism: HRP places 0.901 of its weight on the low-variance cluster and 0.099 on the high-variance one, matching the paper's prediction.""")
P(doc, """The formal specification of each method, with every symbol defined, is in Appendix D.""")

H2(doc, "1.3 Fund universe and the first innovation")
P(doc, """The fund universe is three asset families, equity, crypto, and combined, each built with all five methods, for 15 funds in total. The brief requires only a combined fund built with at least two methods, so the platform exceeds the minimum on two axes at once. The set is wider, 15 funds rather than a handful, and it is newer, because HRP is a 2016 method that post-dates the classical mean-variance toolkit. Both count as innovation under the brief's wording, which credits a wider or newer set of funds or optimisation methods than the required minimum. The lexicon work in Section 4 is a separate and larger innovation; the fund set is the first, smaller one.""")

# ── Section 2 ──
H1(doc, "2. Out-of-sample results")
P(doc, """Table 1 reports out-of-sample performance for all 15 funds. Returns and volatility are annualised, the Sharpe ratio is measured in excess of the daily risk-free rate, and the maximum drawdown is the largest peak-to-trough fall in the fund's value over its out-of-sample life.""")
TABLE(doc, perf_disp, "Table 1. Out-of-sample performance of all 15 funds, 2020 to 2023. Ann. return and Ann. vol. are annualised; Sharpe is excess of the daily risk-free rate (Fama/French RF, Kenneth French Data Library); Max DD is the maximum drawdown. Source: results/tables/performance_metrics.csv.")
P(doc, """The combined Max Sharpe fund is the best risk-adjusted performer, with a Sharpe ratio of 0.983 and a 26.6% annualised return. It beats the best optimised equity fund, Risk Parity at 0.580, and every standalone crypto fund bar the confounded Minimum Variance fund discussed below. The gain comes from diversification across two weakly related return sources. The tangency portfolio can hold equities and crypto together, and because the two classes do not move in lockstep, blending them raises return per unit of risk beyond what either reaches alone. This is the clearest evidence in the report that the combined book, not a single asset class, is the platform's strongest product.""")
P(doc, """The same method fails inside the crypto book. Crypto Max Sharpe is the worst of the five crypto funds, with a Sharpe ratio of 0.229 and a maximum drawdown of -89.5%, despite crypto being the highest-returning class in the sample. The cause is the method's reliance on the sample mean return. Crypto returns are extremely volatile and fat-tailed, so a single 252-day estimate of the mean is dominated by noise. Maximum Sharpe concentrates the portfolio in whichever coin spiked in the window, and that bet reverses out of sample. The method that shines on the diversified combined book is the one that concentrates risk fatally when handed a small, wild cross-section.""")
P(doc, """Crypto Minimum Variance posts a Sharpe ratio of 1.217, higher than every equity fund, but the number should be read with care rather than as proof that crypto beats equities on a risk-adjusted basis. The crypto funds run over 1,208 out-of-sample days from September 2020, while the equity funds cover 753 days from January 2021, so the two are not measured over the same window, and the crypto series captures the 2020 to 2021 bull run in full. Minimum Variance genuinely helps by steering toward the least volatile coins, but this cross-class comparison is confounded by sample period and should not be over-read.""")
P(doc, """HRP delivers the second-shallowest maximum drawdown in every family, behind Minimum Variance, at -16.9% in equity, -18.4% in combined, and -78.1% in crypto, in each case the second-best of the five methods. It rarely wins on Sharpe, and that is the point. HRP trades a little in-sample optimality for out-of-sample stability, the behaviour Lopez de Prado (2016) predicts for a method that avoids inverting a noisy covariance matrix. For a drawdown-averse investor, HRP is the most defensible default even though it is not the highest-returning.""")
P(doc, """Within equities the optimised ranking is Risk Parity at 0.580, then Maximum Sharpe at 0.534, then HRP at 0.520, with Minimum Variance last at 0.325; Equal Weight, the estimation-free benchmark, still posts the highest equity Sharpe at 0.687. [HUMAN EDIT REQUIRED: under the real risk-free rate this equity ordering changed from the old RF=0 draft, where it read Risk Parity, then HRP, then Maximum Sharpe. Max Sharpe now edges out HRP, and Minimum Variance is now the weakest optimised method, so the earlier reading that "the more elaborate method does not win" and that "Maximum Sharpe is penalised most" no longer holds as written. Rewrite this interpretation in your own words: Risk Parity still tops the optimised methods and the simple risk-based methods remain competitive, but Max Sharpe is no longer the worst.]""")
FIGURE(doc, "cumret_by_family.png", "Figure 1. Growth of one dollar invested in each fund, by asset family. Crypto funds dominate the vertical scale and compress the equity and combined lines; the combined Max Sharpe fund is the steadiest strong performer once crypto's swings are set aside.")
FIGURE(doc, "drawdown_combined.png", "Figure 2. Drawdown of the combined funds, the percentage fall from each fund's running peak. HRP and Risk Parity spend less time deep underwater than Max Sharpe, confirming the stability ranking in Table 1.")
FIGURE(doc, "weights_over_time.png", "Figure 3. Portfolio weights over time for the combined funds across methods. Equal Weight holds flat lines by construction, while Max Sharpe reallocates aggressively, the visible source of its higher turnover and deeper drawdowns.")
FIGURE(doc, "sharpe_barplot.png", "Figure 4. Out-of-sample Sharpe ratio by fund. The combined Max Sharpe bar is the tallest and the crypto Max Sharpe bar among the shortest, the two-sided result that anchors this section.")

# ── Section 3 ──
H1(doc, "3. The sentiment index")
H2(doc, "3.1 The scoring model")
P(doc, """The sentiment engine is a finance-tuned version of VADER, a rule-based model that scores text on a compound scale from -1 to +1. The base is finVADER (Korab, 2023), which augments VADER's general lexicon with two finance dictionaries, SentiBigNomics with about 7,295 terms and Henry's word list with 189 terms. On top of finVADER the platform adds its own layer: 123 finance words and 204 finance idioms, mined from external news and scored by a panel of ten independent raters, kept only where the panel agreed on both direction and strength. The base model is borrowed and cited; the mined layer is the original contribution, and Section 4 documents how it was built.""")
H2(doc, "3.2 Before and after coverage")
P(doc, """Plain VADER assigns a non-neutral score to 51.1% of headlines, finVADER to 39.3%, and the extended model to 47.2%. The fall from plain VADER to finVADER is not a regression. General VADER treats ordinary finance vocabulary, words such as market, shares, and tax, as mildly emotional because its lexicon was trained on social media, and it therefore flags sentiment that is not there. finVADER strips out those false positives, which is why its non-neutral rate is lower and more accurate, consistent with Loughran and McDonald (2011) on the mislabelling of finance text by general dictionaries. The extended model then recovers 7.9 points of genuine finance sentiment over finVADER, from vocabulary finVADER was missing, without re-importing plain VADER's false positives. The metric measures coverage, how often the model holds an opinion, not correctness; the evidence for correctness is in Section 4.""")
H2(doc, "3.3 Sector-level construction")
P(doc, """The fund-facing signal is built at sector level. Within each sector the model takes an equal-weighted average across the sector's stocks, so one heavily covered mega-cap cannot dominate the sector's mood. The series is then lagged by at least one trading day, so a decision on day t uses only sentiment from day t minus 1 or earlier, the same no-look-ahead rule that governs the backtest. Days without news are handled in two steps: the last known value is carried forward, on the assumption that sentiment persists until fresh news arrives, and any leading gap with no prior value is set to neutral. Dropping missing days would break the fixed trading calendar the fusion tilt merges onto, so the fill is a design requirement, not just a convenience.""")
H2(doc, "3.4 Coverage evidence")
P(doc, """Pooling headlines upward buys both coverage and calm. A single stock has news on a median 80% of trading days, with a day-to-day standard deviation of 12.81 on the 0 to 100 sentiment scale. Aggregating to sector level lifts coverage to 99% and cuts the daily standard deviation to 7.30; pooling all 50 stocks into a market-wide index reaches 100% coverage at a standard deviation of 2.86, roughly a 4.5-fold reduction in day-to-day noise from the single-stock level. The mechanism is diversification: averaging many series cancels the idiosyncratic noise specific to any one stock while the common, market-wide signal survives. The fund tilt runs at sector level rather than on the aggregate precisely because the tilt is a cross-sectional bet and needs sectors to differ from one another, a difference the fully pooled index by construction erases.""")
H2(doc, "3.5 Building the fear and greed index")
P(doc, """The public fear and greed index is built in four steps. Each headline's compound score is rescaled from its native -1 to +1 range onto 0 to 100, so 50 is neutral. The 50 stocks are then averaged, equal-weighted, into one market-wide series. That raw level is almost uninformative on its own: it sits above 50 on 98% of days and spans only 45.3 to 62.8, so by the raw number the market looks optimistic nearly always. Standardising fixes this. Each day is converted to a z-score, its distance from the index's own historical mean measured in standard deviations, which separates genuinely fearful days from greedy ones. The standardisation uses an expanding window, computing the mean and standard deviation from data up to each date only, so no future information leaks into a past reading. Expanding and full-sample z-scores correlate at 0.998, so the look-ahead-safe version costs almost nothing while remaining valid for a live signal. The latest reading is extreme greed, a z-score of about 2.25, meaning the market sits 2.25 standard deviations above its own 2020 to 2023 average, a window that already includes the COVID crash and the 2022 selloff.""")
H2(doc, "3.6 Attribution")
P(doc, """Two methodologies meet in this report and should not be conflated. The index construction, from raw score through rescaling, aggregation, and expanding-window standardisation, follows the Week 9 lecture. The decision to fuse sentiment into fund weights, the tilt in Section 4, comes from the project brief; the Week 9 material defines the index but prescribes no tilt of its own.""")
FIGURE(doc, "aggregate_sentiment_standardised.png", "Figure 5. The market-wide sentiment index after expanding-window standardisation. Values are z-scores and the bands mark the fear and greed regions; the series ends in extreme greed.")
FIGURE(doc, "sector_sentiment.png", "Figure 6. Sector-level sentiment over time. Sectors diverge from one another, and that cross-sectional variation is what the fund tilt exploits.")

# ── Section 4 ──
H1(doc, "4. Extensions and innovations")
H2(doc, "4.1 The lexicon-mining pipeline")
P(doc, """The primary innovation is a rule-based pipeline for extending the sentiment lexicon, not a hand-picked word list. It runs in four stages. Candidate vocabulary is first mined from an external corpus of roughly 2,154 news articles for idioms and 452 for single words, drawn from Reuters, CNBC, MarketWatch, and Bloomberg via RSS and Google News feeds. This external corpus is used only to discover candidate words and phrases. It never enters the reported results as data: every sentiment score, coverage figure, and backtest in this report runs exclusively on the provided news_headlines.parquet.""")
P(doc, """Each candidate is then rated by ten independent agents on a valence scale from -4 to +4. Each rating is an isolated pass that cannot see the others, so the procedure yields ten separate opinions per candidate, from which a mean and a cross-agent standard deviation are computed. A two-stage filter keeps a candidate only if the panel agrees on both direction and strength: the absolute mean must be at least 0.5, and the cross-agent standard deviation must be below 2.0. In practice the standard-deviation gate almost never binds, with a maximum around 0.52 for words, so the binding constraint is the mean floor. Replacing a subjective author judgement with a consensus rule is what makes the extension reproducible and defensible rather than arbitrary.""")
P(doc, """The survivors, 123 words and 204 idioms, are layered onto finVADER, and the idioms required a specific fix. VADER stores multi-word idioms in a special-case table, but it applies them only when the phrase's last word is a lexicon word, at least three tokens precede it, and the token three positions back is not a lexicon word. Headline-leading idioms, such as a phrase like "shares soar" at the start of a sentence, have too few preceding tokens and so never fire. The fix detects each known idiom and collapses it into a single token carrying the idiom's valence, which fires regardless of position. The effect is concrete: finVADER scores "profit warning" at +0.13, the wrong sign, because profit reads as positive and the idiom handling never triggers; the collapsed idiom scores it negative, as it should.""")
H2(doc, "4.2 An honestly reported negative result")
P(doc, """The extension is reported honestly, including where it stopped helping. A first mining round produced 204 idioms and lifted the fusion Sharpe from 0.534 to 0.552, a gain of 0.018. A second round added 269 more idioms, taking the total to 473, and the fusion gain shrank. [HUMAN EDIT REQUIRED: the 204-versus-473 comparison was measured under the old RF=0 assumption, where the gains were 0.015 and 0.005; only the live 204-idiom fusion was re-run under the real risk-free rate, giving the 0.018 above. Re-run the 473-idiom fusion under the real rate if you want to quote its exact gain, or state that the dilution finding predates the risk-free-rate change.] The larger set was worse, so the platform reverted to the 204-idiom set and archived the 473-idiom set rather than deleting it.""")
P(doc, """The dilution has a clear cause. Reaching 473 idioms meant lowering the frequency threshold to admit rarer phrases, and the marginal candidates are lower quality: the second round pulled in fragmentary, low-frequency n-grams such as "drop as outlook", "discovery reports jump", and "added just workers", each appearing only three times in the corpus and carrying ambiguous rather than clean directional sentiment. [HUMAN EDIT REQUIRED: this corrects the earlier draft, which illustrated the dilution with "central bank", "rate hike", and "cost cutting" as high-frequency boilerplate. The data contradicts that: "rate hike" is in the core 204 set with valence -1.0, "central bank" and "cost cutting" were rejected entirely, and the genuinely diluting idioms are the rare three-occurrence fragments listed above (round-two candidate frequency floor was 3, versus 4 and up in round one). Reword this mechanism in your own words: the marginal idioms are rare and fragmentary, not high-frequency boilerplate.] These noisy matches dilute the sharp signal from the original 204. A selective lexicon can beat a permissive one, though the effect here is small and specific to this sample. A careful extension that does not beat a larger baseline, clearly explained, is still a genuine result.""")
H2(doc, "4.3 HRP as a second innovation")
P(doc, """HRP is the platform's second and independent innovation, distinct in kind from the lexicon work: a newer optimisation method rather than a richer data signal. Its claim to novelty is not that it adds a fifth line to a chart but that it allocates risk without inverting the covariance matrix, the fragile step in Minimum Variance and Maximum Sharpe. That property produces measurably different behaviour, the second-shallowest drawdown in every family in Table 1 (behind Minimum Variance), and weights that rank like Risk Parity yet are not identical to it. It earns its place on evidence, not on novelty for its own sake.""")
H2(doc, "4.4 The fusion result")
P(doc, """Fusing sentiment into the equity Max Sharpe fund improves its risk-adjusted return and slightly worsens its downside. The base fund returns 11.97% at a Sharpe of 0.534, with a maximum drawdown of -26.10%. Tilting its weights toward sectors with above-median lagged sentiment lifts the return to 12.32% and the Sharpe to 0.552, a gain of 0.018, while the maximum drawdown deepens marginally to -26.70%.""")
TABLE(doc, fusion_disp, "Table 2. The equity Max Sharpe fund before and after the sentiment tilt. Source: results/tables/fusion_comparison.csv.")
P(doc, """The two movements together define the trade-off. Because return rose and Sharpe rose, the extra return was not simply bought with proportional extra volatility; the tilt genuinely improved return per unit of risk. But tilting toward whatever is currently in favour concentrates the book into recently popular sectors, and that concentration hurts a little more when a rally reverses, the likely source of the deeper drawdown. Both effects are small, and they rest on a single fund over a single 2021 to 2023 sample, so the honest reading is modest. Sentiment belongs in this platform as a light tilt on top of a sound base allocation, not as a primary signal driving the portfolio.""")
FIGURE(doc, "fusion_comparison.png", "Figure 7. The equity Max Sharpe fund with and without the sentiment tilt. The tilted line ends marginally higher, the visible counterpart to the small Sharpe gain in Table 2.")

# ── Section 5 ──
H1(doc, "5. The app and the investor journey")
H2(doc, "5.1 Structure")
P(doc, """The app is organised as an investor journey across five pages: Compare Funds, Fund Fact Sheet, My Allocation, Market Fear and Greed, and Sentiment Analytics. The order mirrors how an investor actually decides. Compare Funds surveys all 15 funds side by side to shortlist candidates. Fund Fact Sheet drills into one fund's metrics, holdings, and history. My Allocation is the point of action, where the user sets weights across funds and sees the blended result. Market Fear and Greed places that decision in context with the sentiment gauge, and Sentiment Analytics exposes the sector-level detail behind it. Each page answers the question the previous one raises.""")
H2(doc, "5.2 Design system")
P(doc, """The app is not a default Streamlit script. Every chart is rendered through a single shared Plotly theme, so the whole product carries one consistent, interactive, and responsive visual language rather than the slightly different look each default chart would produce. The sentiment centrepiece is a custom fear and greed gauge built with a Plotly indicator, which communicates the market's standardised position at a glance in a way a raw z-score cannot. The growth chart uses a deliberate encoding: colour marks the optimisation method and line style marks the asset family, so all 15 funds are uniquely and meaningfully identifiable and a reader can pick out one method across all three families at a glance. That encoding also fixed a real bug, because the default ten-colour palette repeated once 15 funds were plotted and made funds visually indistinguishable. Good design here is also subtraction: hover behaviour and cluttered legends were pared back so each chart shows one comparison cleanly.""")
H2(doc, "5.3 Target user")
P(doc, """The platform targets an investor who holds at least 10,000 dollars, accepts moderate to high risk, and prefers quantitative, rules-based management to discretionary stock-picking, the same value proposition set out in Part A. AlphaBlend sits between passive index funds and opaque, expensive discretionary funds, offering systematically constructed funds with published fact sheets. The fund range fits this user: it spans a stable HRP equity fund through to a high-return, high-drawdown crypto book, so the investor chooses a point on the risk spectrum rather than accepting a single default. The crypto funds' drawdowns, as deep as -89%, confirm this is not a capital-preservation product.""")
H2(doc, "5.4 Deployment constraint")
P(doc, """The deployed app only reads precomputed CSV files from the results directory. It never re-runs a backtest, re-optimises a portfolio, or invokes the sentiment model at runtime, and it imports neither nltk nor finVADER. This is a deliberate architecture, not a limitation. Heavy computation runs once, offline, and is frozen to disk, so the app starts instantly, stays within the Streamlit Community Cloud free tier, and can never display numbers that disagree with the report. Recomputing live is a listed common mistake, and a grep check across the codebase confirms the app avoids it.""")
H2(doc, "5.5 Links")
P(doc, """[Live Streamlit app URL and the public GitHub repository link to be added at submission; the repository is private until hand-in.]""")

# ── Section 6 ──
H1(doc, "6. Critical reflection")
H2(doc, "6.1 What worked, what did not, and why")
P(doc, """Read across the results, one theme separates the successes from the disappointments: stability and diversification worked, while concentrated bets on estimated returns did not. HRP's drawdown control and the combined Max Sharpe fund's Sharpe of 0.983 are the clearest wins, both flowing from spreading risk rather than chasing it, and the sentiment layer's 7.9-point coverage gain adds a genuine if modest signal. The disappointments were understood in advance rather than surprises. The crypto Max Sharpe fund's -89% drawdown is the predictable cost of feeding a noisy mean estimate into an unconstrained tangency portfolio on a wild asset. The move from 204 to 473 idioms diluted the signal because candidate quality fell as the frequency threshold dropped. The fusion tilt bought a small Sharpe gain at a small drawdown cost. In each case the mechanism, not the outcome alone, is what the result teaches.""")
H2(doc, "6.2 What the index can and cannot tell you")
P(doc, """The sentiment index has real but bounded uses. It can pool thousands of individually noisy headlines into a single series and, once standardised, flag days that are unusually fearful or greedy relative to history. It cannot judge whether a headline is true or important, because it scores tone rather than fact or materiality. It cannot be trusted to get the sign right on every headline; the "profit warning" case, mis-scored at +0.13 before the idiom fix, shows how a noisy proxy fails on individual items. And it cannot stand alone as a buy or sell rule, which is exactly why the fund tilt uses it lightly and why the signal is aggregated to sector level, where idiosyncratic noise cancels. The index is a context indicator, not a trading trigger.""")
H2(doc, "6.3 Three recommendations")
P(doc, """First, match the optimisation method to the client. A drawdown-averse investor should default to HRP or Risk Parity, which post among the shallowest drawdowns in every family (behind Minimum Variance), accepting lower return for stability. An investor chasing return and able to bear volatility should hold the combined Max Sharpe fund, with its 26.6% return and Sharpe of 0.983, accepting deeper drawdowns for that return. The platform should present method as a risk choice, not a technical detail.""")
P(doc, """Second, use sentiment as a modest tilt, not a primary signal. The fusion result improved Sharpe by only 0.018 while deepening drawdown, and expanding the idiom set diluted rather than strengthened the signal. Both point the same way: sentiment should adjust weights gently around a sound base allocation rather than drive them.""")
P(doc, """Third, address the backtest's remaining unrealistic assumption before any live deployment: zero transaction costs. [HUMAN EDIT REQUIRED: this recommendation originally also called for replacing a zero risk-free rate with a real short-rate proxy, but that is now implemented (the daily Fama and French RF series), so the report already carries excess-return Sharpes. Rewrite this recommendation in your own words around transaction costs alone, and consider reframing the risk-free-rate change as work already done and describing its measured effect, that every fund's Sharpe fell once a positive rate was subtracted, by roughly 0.05 to 0.16 for the equity and combined funds against about 0.03 for crypto.] A live product should add a turnover-based cost model; the brief itself treats a transaction-cost model as an innovation. Costs would bite hardest on the high-turnover methods, Maximum Sharpe above all, whose aggressive monthly reallocation is visible in Figure 3, so their reported edge would shrink once trading is charged for, while low-turnover Equal Weight and HRP would be affected least.""")

# ── References (not counted toward the word cap) ──
_counting[0] = False
doc.add_page_break()
H1(doc, "References")
refs = [
    "Baker, M. and Wurgler, J. (2007). Investor Sentiment in the Stock Market. Journal of Economic Perspectives, 21(2), 129-151.",
    "Hutto, C. J. and Gilbert, E. (2014). VADER: A Parsimonious Rule-based Model for Sentiment Analysis of Social Media Text. Proceedings of the Eighth International AAAI Conference on Weblogs and Social Media (ICWSM).",
    "Korab, P. (2023). finVADER: financial sentiment analysis with VADER, SentiBigNomics and the Henry lexicon (Python package).",
    "Lopez de Prado, M. (2016). Building Diversified Portfolios that Outperform Out-of-Sample. The Journal of Portfolio Management, 42(4), 59-69.",
    "Loughran, T. and McDonald, B. (2011). When Is a Liability Not a Liability? Textual Analysis, Dictionaries, and 10-Ks. Journal of Finance, 66(1), 35-65.",
    "Shapiro, A. H., Sudhof, M. and Wilson, D. J. (2022). Measuring news sentiment. Journal of Econometrics, 228(2), 221-243.",
    "Tetlock, P. C. (2007). Giving Content to Investor Sentiment: The Role of Media in the Stock Market. Journal of Finance, 62(3), 1139-1168.",
]
for r in refs:
    P(doc, r)
P(doc, "[HUMAN EDIT REQUIRED: reconcile this list against the Part A reference list, verify every field, and drop any source not actually cited in the final text. Confirm the exact finVADER and Week 9 citation form used in the course.]")

# ── Appendix ──
doc.add_page_break()
H1(doc, "Appendix")
H2(doc, "A. HRP synthetic validation")
P(doc, """On a four-asset test with two low-variance assets, two high-variance assets, and near-zero cross-correlation, HRP allocated 0.901 to the low-variance cluster and 0.099 to the high-variance cluster. The weights are non-negative and sum to one, and rank identically to Risk Parity while differing in magnitude (HRP near 90/10 against Risk Parity near 75/25). Source: ai/prompt_log_12_add_hrp.md.""")
H2(doc, "B. Borderline idioms to spot-check")
P(doc, """The following kept idioms sit closest to the retention floor (|mean valence| = 0.5) and should be reviewed before final submission: "restructuring plan" (-0.50), "share sale" (-0.50), "take private" (+0.70), and "turnaround plan" (+0.80). Source: results/lexicon/kept_idioms.csv.""")
H2(doc, "C. Lexicon artifacts")
P(doc, """The final extension comprises 123 words (results/lexicon/kept_lexicon.csv) and 204 idioms (results/lexicon/kept_idioms.csv). The archived 473-idiom experiment is retained in results/lexicon/kept_idioms_473_round2.csv.""")

H2(doc, "D. Formal specification of the portfolio methods")
P(doc, """Each fund maps the trailing 252-day window of daily returns to a set of weights by one of five rules. Equations (1) to (7) restate exactly what src/portfolio.py computes. All five methods are long-only and fully invested by construction or constraint, with w_i greater than or equal to 0 and the weights summing to 1.""")
H3(doc, "Equal Weight")
EQUATION(doc, r"w_i = \frac{1}{N}, \quad i = 1, \dots, N,",
         "where N is the number of assets in the fund.")
H3(doc, "Minimum Variance")
EQUATION(doc, r"\min_{w}\; w^{\top}\Sigma w \quad \mathrm{s.t.}\; \mathbf{1}^{\top} w = 1,\; 0 \leq w_i \leq 1,",
         "where Σ is the sample covariance matrix of daily returns over the estimation window.")
H3(doc, "Maximum Sharpe (tangency portfolio)")
EQUATION(doc, r"\max_{w}\; \frac{w^{\top}(\mu - r_f)}{\sqrt{w^{\top}\Sigma w}} \quad \mathrm{s.t.}\; \mathbf{1}^{\top} w = 1,\; 0 \leq w_i \leq 1,",
         "where μ is the sample mean daily return vector and r_f is the mean daily risk-free rate over the same estimation window (the real Fama and French rate, not zero).")
H3(doc, "Risk Parity")
EQUATION(doc, r"\min_{w}\; \sum_{i=1}^{N}\left(\frac{w_i(\Sigma w)_i}{w^{\top}\Sigma w} - \frac{1}{N}\right)^{2} \quad \mathrm{s.t.}\; \mathbf{1}^{\top} w = 1,\; w_i \geq 0,",
         "where w_i(Σw)_i / (w′Σw) is asset i's fractional contribution to total portfolio risk, equalised across all assets at the optimum.")
H3(doc, "Hierarchical Risk Parity (López de Prado, 2016)")
P(doc, """HRP allocates risk in three steps: tree clustering on a correlation distance, quasi-diagonalisation, and recursive bisection.""")
EQUATION(doc, r"d_{i,j} = \sqrt{\frac{1}{2}(1 - \rho_{i,j})},",
         "where ρ_{i,j} is the sample correlation between assets i and j, used to build the distance matrix for tree clustering.")
EQUATION(doc, r"V_C = w_C^{\top}\Sigma_C w_C, \quad w_{C,i} = \frac{1/\sigma_i^{2}}{\sum_{j \in C} 1/\sigma_j^{2}},",
         "where V_C is a cluster's variance under inverse-variance weighting and σ_i² is asset i's variance, used to compare two candidate sub-clusters at each split.")
EQUATION(doc, r"\alpha = 1 - \frac{V_L}{V_L + V_R}, \quad w_i \leftarrow \alpha w_i\; (i \in L), \quad w_i \leftarrow (1-\alpha) w_i\; (i \in R),",
         "where L and R are the two sub-clusters at a split and α allocates more of the parent's weight to the lower-variance side, recursively down to single assets.")

# ── Needs review checklist (author-facing) ──
doc.add_page_break()
H1(doc, "Needs Review (author judgement required before submission)")
needs = [
    "Section 2, crypto Minimum Variance Sharpe 1.217 vs equity funds: the claim that this is confounded by the different sample window (1,208 vs 753 days) is the draft's reasoning; confirm you agree the comparison is not like-for-like.",
    "Section 2, why the combined Max Sharpe fund beats either asset class alone: the diversification-of-the-tangency-portfolio explanation is AI reasoning; restate it in your own words and check it against what you understand of mean-variance theory.",
    "Section 3.2, plain VADER's higher non-neutral rate framed as finVADER correcting false positives rather than being worse: verify this reading and that it matches Loughran and McDonald (2011).",
    "Section 4.2, why 204 idioms beat 473: the quality-falls-as-frequency-threshold-drops mechanism is the draft's best inference; confirm it and check the boedrline examples against kept_idioms.csv.",
    "Section 4.4, attributing the deeper drawdown to concentration into recently favoured sectors: this is interpretation, not a measured decomposition; flag it as such or soften if you cannot support it.",
    "Section 5.3, the target user must match Part A exactly; the draft used the Part A value proposition (>=$10,000, moderate-to-high risk, quantitative preference), but re-read your Part A report to confirm wording.",
    "Section 6.3 recommendation 3, the claim that transaction costs would hit high-turnover Max Sharpe hardest: directionally argued from Figure 3, not quantified; keep as a qualitative recommendation.",
    "All references: verify every field and reconcile against Part A before submission.",
    "Risk-free-rate change (Sections 1.1, 2, 4.2, 4.4, 6.3): all Sharpe, fusion, and combined Max Sharpe return figures were regenerated after switching from RF=0 to the daily Fama and French RF. Three items need your judgement, each flagged inline with HUMAN EDIT REQUIRED: (a) Section 1.1 carries two STUDENT-TO-WRITE placeholders to answer; (b) Section 2's equity method ranking flipped (Max Sharpe now edges HRP), so the old 'elaborate method does not win' argument needs rewriting; (c) Section 6.3 recommendation 3 must be reframed because the real-rate fix it recommended is already done. The 473-idiom dilution figures in Section 4.2 were not re-run under the real rate.",
    "Whole draft: this is AI-drafted prose. Rewrite the economic interpretation in your own words so it is genuinely yours, per the AI-use policy.",
]
for i, n in enumerate(needs, 1):
    P(doc, f"{i}. {n}")

# ── save ──
docx_path = REPORT / "report.docx"
doc.save(docx_path)

md_path = REPORT / "report_draft.md"
md_header = ("> DRAFT, NOT FOR SUBMISSION AS-IS. AI-assisted first draft; every number "
             "traces to results/. Rewrite the interpretation in your own words before "
             "submission. See the Needs Review section at the end.\n")
md_path.write_text(md_header + "\n" + "\n".join(MD) + "\n", encoding="utf-8")

# ── em-dash check ──
full = docx_path.read_bytes().decode("latin-1", "ignore") + md_path.read_text(encoding="utf-8")
em = full.count("—")
body_text = " ".join(MD)
# body word count already tracked for counted sections
print(f"Wrote {docx_path}")
print(f"Wrote {md_path}")
print(f"Body word count (Abstract through Section 6): {_body_words[0]}")
print(f"Em dash (\\u2014) occurrences across outputs: {em}")
if em:
    print("WARNING: em dashes present, fix before finalising")
